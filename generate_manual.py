#!/usr/bin/env python3
"""Generate GOIL HSSEQ Portal User Manual as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Colour palette ────────────────────────────────────────────────────────────
ORANGE  = RGBColor(0xF4, 0x79, 0x20)
DARK    = RGBColor(0x0D, 0x11, 0x17)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
GRAY    = RGBColor(0x44, 0x44, 0x44)
LGRAY   = RGBColor(0x99, 0x99, 0x99)
GREEN   = RGBColor(0x3F, 0xB9, 0x50)
RED     = RGBColor(0xDA, 0x36, 0x33)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text.upper())
    set_font(run, size=20, bold=True, color=ORANGE)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'F47920')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=14, bold=True, color=ORANGE)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=GRAY)
    return p

def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=11, color=GRAY)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11, color=GRAY)
    return p

def note(text, kind='note'):
    """kind: 'note' | 'tip' | 'warning'"""
    colors = {'note': (RGBColor(0x2F,0x81,0xF7), '2F81F7'),
              'tip':  (GREEN, '3FB950'),
              'warning': (RED, 'DA3633')}
    rgb, hex_col = colors.get(kind, colors['note'])
    labels = {'note': 'NOTE', 'tip': 'TIP', 'warning': 'WARNING'}

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Inches(0.2)

    # left border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '16')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), hex_col)
    pBdr.append(left)
    pPr.append(pBdr)

    label_run = p.add_run(labels[kind] + ': ')
    set_font(label_run, size=10, bold=True, color=rgb)
    body_run = p.add_run(text)
    set_font(body_run, size=10, italic=True, color=GRAY)
    return p

def step(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    num_run = p.add_run(f'Step {num}. ')
    set_font(num_run, size=11, bold=True, color=ORANGE)
    body_run = p.add_run(text)
    set_font(body_run, size=11, color=GRAY)
    return p

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('GOIL HSSEQ PORTAL')
set_font(r, size=36, bold=True, color=ORANGE)

p2 = doc.add_paragraph()
p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(8)
r2 = p2.add_run('User Manual')
set_font(r2, size=24, bold=False, color=GRAY)

p3 = doc.add_paragraph()
p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(4)
r3 = p3.add_run('Health, Safety, Security, Environment & Quality Department')
set_font(r3, size=13, color=LGRAY)

p4 = doc.add_paragraph()
p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(4)
r4 = p4.add_run('Ghana Oil Company Limited (GOIL)')
set_font(r4, size=13, bold=True, color=GRAY)

doc.add_paragraph()
doc.add_paragraph()

p5 = doc.add_paragraph()
p5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run(f'Version 1.0  |  {datetime.date.today().strftime("%B %Y")}')
set_font(r5, size=10, color=LGRAY)

p6 = doc.add_paragraph()
p6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p6.add_run('Restricted to Authorised GOIL HSSEQ Personnel Only')
set_font(r6, size=9, italic=True, color=RED)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════
heading1('Table of Contents')

toc_entries = [
    ('1', 'Introduction', 4),
    ('2', 'Getting Started – Login & Navigation', 5),
    ('3', 'Home Dashboard (Welcome Screen)', 6),
    ('4', 'Module 1 – Risk Assessment & Facility Inspection', 7),
    ('  4.1', 'Facility Register', 7),
    ('  4.2', 'Conducting an Inspection', 8),
    ('  4.3', 'Risk Scoring', 9),
    ('  4.4', 'Corrective Actions', 10),
    ('  4.5', 'Assessment Dashboard', 10),
    ('5', 'Module 2 – Incident & Hazard Reporting', 11),
    ('6', 'Module 3 – NPA Monitoring Shortcomings', 12),
    ('7', 'Module 4 – Compliance & Regulatory Unit', 14),
    ('8', 'Module 5 – HSSEQ Program & KPI (Coming Soon)', 15),
    ('9', 'User Roles & Permissions', 15),
    ('10', 'Frequently Asked Questions', 16),
    ('11', 'Glossary', 17),
]

for num, title, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r_num = p.add_run(f'{num}  ')
    set_font(r_num, size=11, bold=bool(not num.startswith('  ')), color=ORANGE if not num.startswith('  ') else GRAY)
    r_title = p.add_run(title)
    set_font(r_title, size=11, color=GRAY)
    # tab / dots / page
    r_dots = p.add_run('  ' + ('.' * max(1, 60 - len(num) - len(title))) + f'  {page}')
    set_font(r_dots, size=10, color=LGRAY)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
heading1('1. Introduction')

body(
    'The GOIL HSSEQ Portal is a web-based management information system designed exclusively for the '
    'Health, Safety, Security, Environment & Quality (HSSEQ) Department of Ghana Oil Company Limited (GOIL). '
    'It brings together five functional modules under a single, secure platform to support GOIL\u2019s commitment '
    'to operational safety, regulatory compliance, and continuous improvement.'
)

heading2('1.1 Purpose of This Manual')
body('This manual guides authorised GOIL HSSEQ personnel through every feature of the portal. '
     'It covers logging in, navigating the interface, conducting facility inspections, recording incidents, '
     'managing NPA shortcomings, tracking compliance obligations, and monitoring KPIs.')

heading2('1.2 Scope')
body('This manual covers the following portal modules:')
for m in [
    'Risk Assessment & Facility Inspection',
    'Incident & Hazard Reporting',
    'NPA Monitoring Shortcomings',
    'Compliance & Regulatory Unit',
    'HSSEQ Program & KPI (planned)',
]:
    bullet(m)

heading2('1.3 System Requirements')
body('The portal runs entirely in a web browser. No software installation is required.')
for r in [
    'Modern browser: Google Chrome 110+, Microsoft Edge 110+, or Mozilla Firefox 110+',
    'Stable internet connection',
    'Screen resolution: 1280 × 720 minimum (1920 × 1080 recommended)',
    'Authorised GOIL email address and password',
]:
    bullet(r)

note('The portal is mobile-responsive and can be used on tablets and smartphones, though a desktop or '
     'laptop is recommended for inspection data entry.', 'tip')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. GETTING STARTED
# ═══════════════════════════════════════════════════════════════════════════════
heading1('2. Getting Started – Login & Navigation')

heading2('2.1 Accessing the Portal')
body('Open your web browser and navigate to the GOIL HSSEQ Portal URL provided by your system administrator. '
     'The login page will appear, featuring the GOIL logo on the left panel and a sign-in form on the right.')

heading2('2.2 Signing In')
step(1, 'Enter your assigned GOIL email address (e.g. name@goil.com.gh) in the Email field.')
step(2, 'Enter your password. Click the eye icon to reveal or hide the password.')
step(3, 'Tick "Keep me signed in" if you are on a trusted device and wish to remain logged in.')
step(4, 'Click Login. The portal will validate your credentials and redirect you to the Home Dashboard.')

note('If your account requires a password change on first login, a "Set New Password" dialog will appear automatically. '
     'Enter a new password of at least 8 characters and confirm it to continue.', 'warning')

heading2('2.3 Forgot Password')
step(1, 'Click the Forgot password? link on the sign-in form.')
step(2, 'Enter your GOIL email address and click Send Reset Link.')
step(3, 'Check your inbox for a password reset email. Click the link in the email to set a new password.')
step(4, 'After resetting, return to the portal and sign in with your new password.')

heading2('2.4 Navigating the Portal')
body('After signing in you will see the Home Dashboard. The interface has two main navigation elements:')

heading3('Sidebar (left panel)')
for item in [
    'GOIL "G" logo — click to return to the Home Dashboard at any time.',
    'Navigation items — each module is listed with an icon and label.',
    'Collapse button — click the small circle arrow on the sidebar edge to collapse it to icon-only mode, saving screen space.',
    'User section (bottom of sidebar) — displays your name, role, and a menu for account options.',
]:
    bullet(item)

heading3('Top Bar')
for item in [
    'Breadcrumb trail — shows your current location within the portal.',
    'User pill — your name and avatar; click to access account settings.',
    'Mobile menu button (on small screens) — opens the sidebar overlay.',
]:
    bullet(item)

heading2('2.5 Signing Out')
body('Click your name/avatar in the sidebar footer or top bar, then select Sign Out (or Logout). '
     'You will be returned to the login page. Always sign out when using a shared device.')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. HOME DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════════
heading1('3. Home Dashboard (Welcome Screen)')

body('The Home Dashboard is the first screen you see after logging in. It provides a high-level overview '
     'of activity across all modules and quick-access cards to each module.')

heading2('3.1 Summary Widgets')
body('Four live summary widgets appear at the top of the dashboard:')

for widget, desc in [
    ('Open Draft Assessments', 'Number of facility inspections that have been started but not yet submitted.'),
    ('Overdue Actions', 'Corrective actions whose due dates have passed without being closed.'),
    ('Critical Findings', 'Findings scored as Critical risk that require immediate attention.'),
    ('Due Assessments', 'Facilities whose next scheduled inspection is due within 30 days.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    r_w = p.add_run(f'{widget}: ')
    set_font(r_w, size=11, bold=True, color=ORANGE)
    r_d = p.add_run(desc)
    set_font(r_d, size=11, color=GRAY)

heading2('3.2 Module Navigation Cards')
body('Below the widgets are module cards. Click any card to open that module. '
     'Cards with a "Coming Soon" badge are planned but not yet available.')

heading2('3.3 Dark / Light Mode')
body('A sun/moon icon in the top-right corner of the login page, and within the portal, allows you to '
     'switch between dark mode (default) and light mode. Your preference is saved automatically.')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. MODULE 1 – RISK ASSESSMENT & FACILITY INSPECTION
# ═══════════════════════════════════════════════════════════════════════════════
heading1('4. Module 1 – Risk Assessment & Facility Inspection')

body('This is the primary operational module. It supports area-based physical site inspections, '
     'automatic risk generation from deviations, and corrective action management.')

heading2('4.1 Facility Register')

body('The Facility Register lists all GOIL facilities under the HSSEQ inspection programme. '
     'Access it from the sidebar under Risk Assessment → Facilities.')

heading3('Facility Status Chips')
body('Each facility card displays a colour-coded due-state chip:')
for chip, meaning in [
    ('On Track (green)', 'Next inspection is more than 30 days away.'),
    ('Due Soon (amber)', 'Next inspection is within the next 30 days.'),
    ('Due Today (orange)', 'Inspection is due today.'),
    ('Overdue (red)', 'Inspection due date has passed.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    r_c = p.add_run(f'• {chip}: ')
    set_font(r_c, size=11, bold=True, color=GRAY)
    r_m = p.add_run(meaning)
    set_font(r_m, size=11, color=GRAY)

heading3('Filtering Facilities')
body('Use the filter bar to narrow facilities by region, type, or due status. '
     'Click Reset Filters to clear all selections.')

heading3('Starting or Resuming an Inspection')
step(1, 'Click a facility card to open its detail view.')
step(2, 'Click Start Inspection to begin a new assessment, or Resume to continue a saved draft.')
note('Only one draft assessment can be open per facility at a time. '
     'Submit or discard the existing draft before starting a new one.', 'note')

heading2('4.2 Conducting an Inspection')

heading3('Area Navigator')
body('The inspection is divided into physical site areas (e.g. Forecourt, Canopy, Office). '
     'The area navigator on the left shows your progress through each area with completion percentages.')

heading3('Answering Checklist Questions')
body('For each question select one of four responses:')
for resp, meaning in [
    ('Compliant', 'The item fully meets the standard. No action required.'),
    ('Partially Compliant', 'The item partially meets the standard. A finding will be raised.'),
    ('Non-Compliant', 'The item fails to meet the standard. A finding and risk card will be raised.'),
    ('Not Applicable', 'The item does not apply to this facility or area.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    r_c = p.add_run(f'• {resp}: ')
    set_font(r_c, size=11, bold=True, color=GRAY)
    r_m = p.add_run(meaning)
    set_font(r_m, size=11, color=GRAY)

heading3('Mandatory Fields for Deviations')
body('When you select Partially Compliant or Non-Compliant, the following fields become mandatory:')
for f in ['Observation / Comment — describe what was observed.',
          'Evidence — attach a photo or enter a description of evidence.']:
    bullet(f)

body('For Non-Compliant responses on critical items, three additional fields are required:')
for f in ['Responsible Owner — the person accountable for resolving the issue.',
          'Due Date — the target date for corrective action completion.',
          'Containment Action — the immediate action taken to mitigate risk.']:
    bullet(f)

note('You cannot submit the inspection until all mandatory fields are completed. '
     'The system highlights incomplete items in red.', 'warning')

heading3('Saving Progress')
body('The inspection auto-saves as a draft whenever you navigate between areas. '
     'You can close the browser and return later; your progress will be preserved.')

heading3('Submitting the Inspection')
step(1, 'Navigate to the Review & Submit screen (last step in the area navigator).')
step(2, 'Review the completion summary. All areas must reach 100% to enable submission.')
step(3, 'Click Submit Assessment. The assessment is locked and a summary report is generated.')

heading2('4.3 Risk Scoring')

body('The portal automatically calculates risk for every Non-Compliant or Partially Compliant finding '
     'using a 5 × 5 matrix (Severity × Likelihood).')

heading3('Risk Bands')
for band, range_, color in [
    ('Low',      '1–4',   'Green'),
    ('Medium',   '5–9',   'Amber'),
    ('High',     '10–16', 'Orange'),
    ('Critical', '17–25', 'Red'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    r_b = p.add_run(f'• {band} ({range_} — {color}): ')
    set_font(r_b, size=11, bold=True, color=GRAY)

body('Compliance scoring (percentage) is calculated separately from risk scoring. '
     'A facility can have a high compliance score but still have critical risk items.')

heading2('4.4 Corrective Actions')

body('Every finding automatically generates a corrective action card. '
     'Access corrective actions from Risk Assessment → Actions in the sidebar.')

heading3('Action Statuses')
for status, meaning in [
    ('Open', 'Action has been raised and is awaiting attention.'),
    ('In Progress', 'Work has started on the corrective action.'),
    ('Closed', 'The action has been completed and verified.'),
    ('Overdue', 'The due date has passed and the action is not yet closed.'),
]:
    bullet(f'{status}: {meaning}')

heading3('Updating an Action')
step(1, 'Open the action card from the Actions list.')
step(2, 'Add a progress note or update the status.')
step(3, 'If closing, provide closure evidence and confirmation.')
step(4, 'Click Save. The action history log is updated automatically.')

heading2('4.5 Assessment Dashboard')

body('The Assessment Dashboard provides a visual summary of all submitted assessments.')

heading3('Key Panels')
for panel in [
    'Compliance Score Trend — line chart of average compliance over time.',
    'Risk Distribution — breakdown of findings by risk band.',
    'Category / Area View — compliance scores by inspection area and category.',
    'Action Status Summary — open, in-progress, closed, and overdue action counts.',
]:
    bullet(panel)

heading3('Filters')
body('Use the filter bar at the top to narrow results by facility, date range, or assessment type. '
     'Click Reset to return to the default view.')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. MODULE 2 – INCIDENT & HAZARD REPORTING
# ═══════════════════════════════════════════════════════════════════════════════
heading1('5. Module 2 – Incident & Hazard Reporting')

body('This module provides a structured form for reporting workplace incidents, near-misses, and hazards '
     'across all GOIL facilities. Submitted reports are routed for review, investigation, and closure.')

heading2('5.1 Submitting a New Report')
step(1, 'In the sidebar, navigate to Incident Reporting → New Report.')
step(2, 'Select the report type: Incident, Near Miss, or Hazard.')
step(3, 'Fill in the Incident Details section:')
for f in [
    'Date and Time of occurrence',
    'Facility / Location',
    'Brief Description of what happened',
    'Immediate actions taken',
]:
    bullet(f, level=1)
step(4, 'Complete the People Involved section (if applicable).')
step(5, 'Attach supporting evidence (photographs, witness statements).')
step(6, 'Click Submit Report. A reference number is assigned.')

note('Incidents involving injury, fatality, or significant property damage must be reported within 24 hours '
     'in accordance with GOIL HSSEQ policy and applicable regulations.', 'warning')

heading2('5.2 Incident List')
body('The Incident List (sidebar: Incident Reporting → All Reports) shows all submitted reports with '
     'filters for status, date range, type, and facility. Click any row to open the report detail.')

heading2('5.3 Incident Dashboard')
body('The Incident Dashboard (sidebar: Incident Reporting → Dashboard) provides statistical summaries:')
for item in [
    'Incident counts by type and month',
    'Severity distribution',
    'Facility breakdown',
    'Open investigation status',
]:
    bullet(item)

heading2('5.4 Incident Archive')
body('Closed and archived reports are accessible from Incident Reporting → Archive. '
     'Archived reports are read-only and cannot be re-opened without administrator action.')

heading2('5.5 Admin Functions')
body('Users with the HSSEQ Admin role can access Incident Reporting → Admin to:')
for item in [
    'Assign investigations to team members',
    'Change report status',
    'Add internal investigation notes',
    'Export reports to PDF or CSV',
]:
    bullet(item)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. MODULE 3 – NPA MONITORING SHORTCOMINGS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('6. Module 3 – NPA Monitoring Shortcomings')

body('This module tracks shortcomings identified during National Petroleum Authority (NPA) regulatory '
     'inspections. It provides a workspace to log, track, and close NPA findings before or after the '
     'NPA follow-up visit.')

heading2('6.1 NPA Workspace Overview')

body('The NPA Workspace is the central hub for this module. It is reached via Compliance → NPA Shortcomings '
     'in the sidebar. The workspace has three tabs:')
for tab, desc in [
    ('Overview / Dashboard', 'Summary statistics: total shortcomings, open items, overdue items, and closure rate.'),
    ('Records', 'Full list of all NPA shortcomings with search, filters, and bulk actions.'),
    ('Add Shortcoming', 'Form to log a new NPA finding.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    r_t = p.add_run(f'• {tab}: ')
    set_font(r_t, size=11, bold=True, color=ORANGE)
    r_d = p.add_run(desc)
    set_font(r_d, size=11, color=GRAY)

heading2('6.2 Logging a New NPA Shortcoming')
step(1, 'Navigate to the Add Shortcoming tab (or click the + New Shortcoming button).')
step(2, 'Complete the shortcoming details:')
for f in [
    'NPA Inspection Date — date of the NPA site visit.',
    'Facility — the facility where the shortcoming was identified.',
    'Shortcoming Category — e.g. Storage, Safety Equipment, Housekeeping.',
    'Description — full description of the finding as stated by the NPA inspector.',
    'NPA Severity — classification assigned by the NPA (e.g. Major, Minor, Observation).',
    'Responsible Officer — the GOIL staff member accountable for resolution.',
    'Target Closure Date — the deadline for resolving the shortcoming.',
]:
    bullet(f, level=1)
step(3, 'Optionally attach supporting documentation.')
step(4, 'Click Save Record. The shortcoming appears in the Records list.')

heading2('6.3 Tracking and Updating Shortcomings')

heading3('Record Statuses')
for status, meaning in [
    ('Open', 'Shortcoming has been logged; no corrective action yet.'),
    ('In Progress', 'Corrective work is underway.'),
    ('Resolved – Pending NPA Confirmation', 'GOIL has resolved the issue; awaiting NPA sign-off.'),
    ('Closed', 'NPA has confirmed the shortcoming is resolved.'),
    ('Overdue', 'The target closure date has passed without closure.'),
]:
    bullet(f'{status}: {meaning}')

heading3('Updating a Record')
step(1, 'In the Records list, click the row or the View/Edit button for the shortcoming.')
step(2, 'On the Record Detail page, scroll to the Activity Log section.')
step(3, 'Add a progress note describing actions taken.')
step(4, 'Update the Status field as appropriate.')
step(5, 'Attach closure evidence if marking as Resolved or Closed.')
step(6, 'Click Save Changes.')

heading2('6.4 Searching and Filtering Records')
body('In the Records tab, use the search bar to find shortcomings by keyword, facility, or NPA reference number. '
     'Apply filters for:')
for f in ['Status', 'Facility', 'Shortcoming Category', 'Date Range', 'Responsible Officer']:
    bullet(f)
body('Click Export to download the filtered list as a CSV file for reporting.')

heading2('6.5 NPA Dashboard Statistics')
body('The Overview tab displays four key metrics:')
for metric in [
    'Total Shortcomings — all shortcomings ever logged.',
    'Open Items — shortcomings not yet resolved.',
    'Overdue Items — shortcomings past their target closure date.',
    'Closure Rate — percentage of shortcomings closed.',
]:
    bullet(metric)

note('The NPA module is scoped per user for non-admin roles — you will only see records assigned to '
     'or created by your account unless you have Admin privileges.', 'note')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. MODULE 4 – COMPLIANCE & REGULATORY UNIT
# ═══════════════════════════════════════════════════════════════════════════════
heading1('7. Module 4 – Compliance & Regulatory Unit')

body('The Compliance & Regulatory Unit module provides a central register of GOIL’s regulatory and '
     'internal compliance obligations across all HSSEQ disciplines: HSE, Quality, ISO, Security, and Common.')

heading2('7.1 Unit Tabs')
body('The module is organised into five unit tabs accessible from the top navigation bar:')
for tab, desc in [
    ('HSE', 'Health, Safety & Environment regulatory requirements and permit obligations.'),
    ('Quality', 'Product quality standards, laboratory compliance, and QMS requirements.'),
    ('ISO', 'ISO certification requirements (ISO 9001, ISO 14001, ISO 45001 etc.).'),
    ('Security', 'Security-related compliance obligations and access control standards.'),
    ('Common', 'Obligations shared across multiple units, e.g. fire safety, first aid.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    r_t = p.add_run(f'• {tab}: ')
    set_font(r_t, size=11, bold=True, color=ORANGE)
    r_d = p.add_run(desc)
    set_font(r_d, size=11, color=GRAY)

heading2('7.2 Viewing Compliance Items')
body('Each unit tab shows a list of compliance obligations. Click any item to expand its detail, which includes:')
for f in [
    'Regulatory Reference — the specific clause, permit, or standard.',
    'Obligation Description — what GOIL must do to comply.',
    'Responsible Unit — the team accountable for this obligation.',
    'Review Frequency — how often compliance must be verified.',
    'Status — Compliant, Partial, Non-Compliant, or Pending Review.',
    'Next Review Date.',
    'Evidence / Notes.',
]:
    bullet(f)

heading2('7.3 Updating Compliance Status')
step(1, 'Open the compliance item by clicking it in the list.')
step(2, 'Click Edit or the status dropdown to change the compliance status.')
step(3, 'Add supporting notes or attach evidence.')
step(4, 'Set the Next Review Date.')
step(5, 'Click Save.')

note('Only users with the Compliance Officer or HSSEQ Admin role can edit compliance item statuses.', 'note')

heading2('7.4 AMC Action Tracking')
body('The Compliance module integrates with the Action Management Centre (AMC). '
     'Non-compliant items automatically generate actions visible in the AMC folder counts shown in the sidebar.')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. MODULE 5 – HSSEQ PROGRAM & KPI
# ═══════════════════════════════════════════════════════════════════════════════
heading1('8. Module 5 – HSSEQ Program & KPI (Coming Soon)')

body('The HSSEQ Program & KPI module is planned for a future release. When available, it will include:')
for item in [
    'HSSEQ programme targets and baselines.',
    'KPI dashboards with trend analysis.',
    'Assessment coverage tracking (% of facilities assessed within the cycle window).',
    'Action closure discipline reporting.',
    'Repeated-issues and chronic-recurrence monitoring.',
    'Management review packs and export.',
]:
    bullet(item)

note('This module is currently visible in the portal but shows a "Coming Soon" placeholder. '
     'No data entry is available until the module is activated.', 'note')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. USER ROLES & PERMISSIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('9. User Roles & Permissions')

body('Access to portal features is controlled by user roles. Your role is assigned by the system administrator '
     'and is displayed in the sidebar footer.')

heading2('9.1 Role Summary')

roles = [
    ('Inspector / Field Officer',
     'Conduct inspections, submit reports, log NPA shortcomings, view own records.'),
    ('HSSEQ Officer',
     'All Inspector permissions plus: view all records, update corrective actions, edit compliance items.'),
    ('Compliance Officer',
     'All Officer permissions plus: manage compliance register, assign obligations.'),
    ('HSSEQ Manager / Admin',
     'Full access to all modules, all records, admin functions, user management, and data exports.'),
]

for role, perms in roles:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.15)
    r_role = p.add_run(role)
    set_font(r_role, size=12, bold=True, color=ORANGE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.3)
    r_p = p2.add_run(perms)
    set_font(r_p, size=11, color=GRAY)

heading2('9.2 Requesting Access Changes')
body('To request a role change or new user account, contact your HSSEQ system administrator '
     'or send an email to the GOIL ICT Helpdesk with your name, department, and required access level.')

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. FREQUENTLY ASKED QUESTIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('10. Frequently Asked Questions')

faqs = [
    ('I cannot log in. What should I do?',
     'Ensure you are using your GOIL email address and the correct password. '
     'If you have forgotten your password, click "Forgot password?" on the login page. '
     'If the problem persists, contact your system administrator.'),
    ('My inspection draft has disappeared. Is it lost?',
     'Drafts are stored in your browser. If you cleared your browser data or switched browsers, '
     'the draft may no longer be available. Always submit assessments promptly to avoid data loss.'),
    ('Can I edit a submitted assessment?',
     'Submitted assessments are locked and cannot be directly edited. Contact an HSSEQ Admin '
     'to request a correction or amendment.'),
    ('Why are some dashboard counts different from what I expect?',
     'Non-admin users see counts scoped to their own records only. Admin users see counts for all users. '
     'Contact your administrator if figures appear inconsistent.'),
    ('Can I use the portal offline?',
     'The portal requires an internet connection for most functions. Certain browsers may cache the '
     'pages but data submission requires connectivity.'),
    ('How do I export data?',
     'Most list views have an Export button that downloads the current filtered view as a CSV file. '
     'Completed assessments can be printed or exported to JSON from the Review & Submit screen.'),
    ('Who do I contact for technical support?',
     'For technical issues, contact the GOIL ICT Helpdesk or your designated HSSEQ system administrator.'),
]

for q, a in faqs:
    heading3('Q: ' + q)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    r_label = p.add_run('A: ')
    set_font(r_label, size=11, bold=True, color=ORANGE)
    r_answer = p.add_run(a)
    set_font(r_answer, size=11, color=GRAY)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. GLOSSARY
# ═══════════════════════════════════════════════════════════════════════════════
heading1('11. Glossary')

terms = [
    ('AMC', 'Action Management Centre — the system that tracks corrective and preventive actions.'),
    ('Assessment', 'A structured facility inspection conducted using the GOIL HSSEQ checklist.'),
    ('Corrective Action (CA)', 'A planned activity to resolve a finding and prevent recurrence.'),
    ('Critical Finding', 'A finding scored 17–25 on the 5×5 risk matrix, requiring immediate action.'),
    ('Deviation', 'A Non-Compliant or Partially Compliant response that triggers a finding.'),
    ('Finding', 'An identified non-conformance or area of concern raised during an inspection.'),
    ('GOIL', 'Ghana Oil Company Limited.'),
    ('HSSEQ', 'Health, Safety, Security, Environment & Quality.'),
    ('ISO', 'International Organisation for Standardisation.'),
    ('KPI', 'Key Performance Indicator — a measurable value that demonstrates effectiveness.'),
    ('NPA', 'National Petroleum Authority — the Ghanaian regulatory body for petroleum products.'),
    ('Risk Band', 'A category (Low, Medium, High, Critical) derived from the risk score.'),
    ('Risk Score', 'Severity × Likelihood on a 5×5 matrix, ranging from 1 to 25.'),
    ('Shortcoming', 'A finding raised by the NPA during a regulatory inspection.'),
    ('Template', 'The structured question set used for a specific facility type (e.g. Fuel Station).'),
]

for term, definition in terms:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r_term = p.add_run(f'{term}:  ')
    set_font(r_term, size=11, bold=True, color=ORANGE)
    r_def = p.add_run(definition)
    set_font(r_def, size=11, color=GRAY)

# ── Footer note ───────────────────────────────────────────────────────────────
doc.add_paragraph()
p_foot = doc.add_paragraph()
p_foot.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_foot = p_foot.add_run(
    f'GOIL HSSEQ Portal — User Manual v1.0  |  {datetime.date.today().strftime("%B %Y")}  |  '
    'Restricted to Authorised GOIL HSSEQ Personnel Only'
)
set_font(r_foot, size=9, italic=True, color=LGRAY)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/Users/macprom4/Documents/Claude project 2/goil-hsseq-app/GOIL_HSSEQ_Portal_User_Manual.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
